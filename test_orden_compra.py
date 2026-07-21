"""
Caso de prueba del LECTOR DE PO (orden_compra.py).

Cubre lo que se puede validar sin la key real de Anthropic (que solo vive en Railway):
  1) EXTRACCIÓN determinista de texto+tablas en los 4 formatos (PDF, Excel, Word, correo .eml).
  2) NORMALIZACIÓN: se simula (stub) la respuesta del LLM para validar la lógica robusta —
     limpieza de ```-fences, defaults del esquema, y fallback ante JSON inválido / doc vacío.

Correr:  .venv/bin/python test_orden_compra.py
(reportlab solo se usa aquí para generar un PDF de muestra; NO es dependencia de producción.)
"""

import io
import json

import orden_compra as oc

# Datos "verdad" que deben aparecer tras extraer/normalizar cualquier formato.
CLIENTE = "Aceros del Norte SA de CV"
PO_NUM = "4501099887"
ITEMS = [("CAT.632-DA", "Pedal Linemaster", 5, "385.00"),
         ("A-81359", "Disco Makita 4-3/8", 2, "712.14")]

_fallos: list[str] = []


def check(nombre: str, cond: bool, detalle: str = ""):
    print(f"  {'✅' if cond else '❌'} {nombre}" + (f" — {detalle}" if detalle and not cond else ""))
    if not cond:
        _fallos.append(nombre)


# ── Generadores de PO de muestra ──────────────────────────────
def po_xlsx() -> bytes:
    import openpyxl
    wb = openpyxl.Workbook(); ws = wb.active; ws.title = "PO"
    ws.append(["ORDEN DE COMPRA"])
    ws.append(["Cliente:", CLIENTE]); ws.append(["PO No.:", PO_NUM]); ws.append(["Moneda:", "MXN"])
    ws.append(["Parte", "Descripcion", "Cantidad", "Precio Unit"])
    for pn, desc, q, p in ITEMS:
        ws.append([pn, desc, q, float(p)])
    b = io.BytesIO(); wb.save(b); return b.getvalue()


def po_docx() -> bytes:
    import docx
    d = docx.Document()
    d.add_paragraph(f"ORDEN DE COMPRA - {CLIENTE}")
    d.add_paragraph(f"PO No.: {PO_NUM}   Moneda: MXN")
    tb = d.add_table(rows=1, cols=4); h = tb.rows[0].cells
    h[0].text, h[1].text, h[2].text, h[3].text = "Parte", "Descripcion", "Cant", "Precio"
    for pn, desc, q, p in ITEMS:
        r = tb.add_row().cells
        r[0].text, r[1].text, r[2].text, r[3].text = pn, desc, str(q), p
    b = io.BytesIO(); d.save(b); return b.getvalue()


def po_pdf() -> bytes:
    from reportlab.pdfgen import canvas
    b = io.BytesIO(); c = canvas.Canvas(b)
    lineas = [f"ORDEN DE COMPRA - {CLIENTE}", f"PO No.: {PO_NUM}   Moneda: MXN", ""]
    lineas += [f"{pn}   {desc}   {q}   {p}" for pn, desc, q, p in ITEMS]
    for i, l in enumerate(lineas):
        c.drawString(50, 800 - i * 22, l)
    c.save(); return b.getvalue()


def po_eml() -> bytes:
    cuerpo = (f"Buen dia, adjunto nuestra orden de compra {PO_NUM}.\n"
              f"Empresa: {CLIENTE}\n" + "\n".join(f"{pn} - {desc} - {q} pzas - ${p}" for pn, desc, q, p in ITEMS))
    raw = (f"From: compras@acerosdelnorte.com\r\nSubject: Orden de compra {PO_NUM}\r\n"
           f"Content-Type: text/plain; charset=utf-8\r\n\r\n{cuerpo}")
    return raw.encode("utf-8")


# ── 1) EXTRACCIÓN ─────────────────────────────────────────────
def test_extraccion():
    print("\n[1] EXTRACCIÓN determinista (sin LLM):")
    casos = [("orden.pdf", po_pdf()), ("orden.xlsx", po_xlsx()),
             ("orden.docx", po_docx()), ("orden.eml", po_eml())]
    for nombre, data in casos:
        texto, fmt = oc.extraer_texto(data, nombre)
        base = nombre.split(".")[0]
        ok_pn = all(pn in texto for pn, *_ in ITEMS)
        ok_po = PO_NUM in texto
        ok_cli = "Aceros del Norte" in texto
        check(f"{fmt}: extrae texto", bool(texto.strip()))
        check(f"{fmt}: contiene ambos part numbers", ok_pn, texto[:150])
        check(f"{fmt}: contiene PO number", ok_po)
        check(f"{fmt}: contiene cliente", ok_cli)


# ── 2) NORMALIZACIÓN (LLM simulado) ───────────────────────────
def test_normalizacion(monkeypatch_target=oc):
    print("\n[2] NORMALIZACIÓN (respuesta del LLM simulada):")

    class _U: input_tokens = 0; output_tokens = 0
    class _B:
        def __init__(self, t): self.text = t
    class _R:
        def __init__(self, t): self.content = [_B(t)]; self.usage = _U()

    respuesta_buena = json.dumps({
        "cliente": CLIENTE, "po_number": PO_NUM, "moneda": "MXN",
        "items": [{"part_number": pn, "descripcion": d, "cantidad": q,
                   "precio_unitario": float(p)} for pn, d, q, p in ITEMS],
        "notas": "",
    })

    # Caso A: LLM responde JSON limpio → normaliza bien.
    oc.llm.complete = lambda **kw: _R(respuesta_buena)
    r = oc.normalizar("texto crudo del PO...")
    check("normaliza cliente", r.get("cliente") == CLIENTE)
    check("normaliza po_number", r.get("po_number") == PO_NUM)
    check("normaliza 2 items", len(r.get("items", [])) == 2)
    check("item trae part_number+precio", r["items"][0].get("part_number") == "CAT.632-DA"
          and r["items"][0].get("precio_unitario") == 385.00)

    # Caso B: LLM envuelve en ```json ... ``` → se limpia igual.
    oc.llm.complete = lambda **kw: _R("```json\n" + respuesta_buena + "\n```")
    r2 = oc.normalizar("texto")
    check("limpia cercas de código ```", r2.get("po_number") == PO_NUM and len(r2.get("items", [])) == 2)

    # Caso C: LLM devuelve basura → fallback controlado, no revienta.
    oc.llm.complete = lambda **kw: _R("lo siento, no pude")
    r3 = oc.normalizar("texto")
    check("fallback ante JSON inválido", "error" in r3 and r3.get("items") == [])

    # Caso D: documento vacío → error temprano sin llamar al LLM.
    r4 = oc.normalizar("   ")
    check("documento vacío → error temprano", "error" in r4)


if __name__ == "__main__":
    print("=" * 60)
    print("CASO DE PRUEBA — LECTOR DE PO (orden_compra.py)")
    print("=" * 60)
    test_extraccion()
    test_normalizacion()
    print("\n" + "=" * 60)
    if _fallos:
        print(f"❌ {len(_fallos)} fallo(s): {_fallos}")
        raise SystemExit(1)
    print("✅ TODO VERDE — extracción y normalización funcionan")
