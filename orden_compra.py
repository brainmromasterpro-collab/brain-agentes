"""
LECTOR DE ÓRDENES DE COMPRA (PO) — Fase 1 del stream Sales Order
================================================================
Lee la orden de compra que manda el cliente (PDF / Excel / Word / correo), saca el TEXTO de forma
determinista (sin costo de visión) y luego un LLM chico lo NORMALIZA a datos estructurados:

    {cliente, po_number, moneda, items: [{part_number, descripcion, cantidad, precio_unitario}]}

Regla de oro (igual que en búsquedas): NO inventar. Si un dato no está claro, se deja vacío y se
marca; es mejor "no lo encontré" que un número equivocado que genere una orden de venta errónea.

Esta capa SOLO lee y estructura. El cotejo contra cotizaciones y la escritura al CRM viven en
agente_chat.py y NO tocan nada aquí.
"""

import io
import os
import json
import logging

import httpx

import llm

log = logging.getLogger("orden_compra")

# Modelo para normalizar (barato). Configurable por si se quiere mover a OpenRouter/DeepSeek luego.
PO_MODEL = os.environ.get("PO_MODEL", "claude-haiku-4-5-20251001")


# ─────────────────────────────────────────────────────────────
# 1. EXTRACCIÓN DE TEXTO (determinista, por formato)
# ─────────────────────────────────────────────────────────────
def _ext(nombre: str, mime: str = "") -> str:
    n = (nombre or "").lower()
    for e in ("pdf", "xlsx", "xls", "docx", "doc", "csv", "txt", "eml"):
        if n.endswith("." + e):
            return e
    m = (mime or "").lower()
    if "pdf" in m: return "pdf"
    if "sheet" in m or "excel" in m: return "xlsx"
    if "word" in m or "document" in m: return "docx"
    if "csv" in m: return "csv"
    return "txt"


def _texto_pdf(data: bytes) -> str:
    import pdfplumber
    partes: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        for pg in pdf.pages:
            t = pg.extract_text() or ""
            if t:
                partes.append(t)
            # Las tablas son clave en un PO (renglones de producto): añadirlas como texto tabulado.
            for tabla in (pg.extract_tables() or []):
                for fila in tabla:
                    celdas = [str(c).strip() for c in fila if c is not None]
                    if celdas:
                        partes.append(" | ".join(celdas))
    return "\n".join(partes)


def _texto_xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    partes: list[str] = []
    for ws in wb.worksheets:
        partes.append(f"[Hoja: {ws.title}]")
        for fila in ws.iter_rows(values_only=True):
            celdas = [str(c).strip() for c in fila if c is not None and str(c).strip()]
            if celdas:
                partes.append(" | ".join(celdas))
    return "\n".join(partes)


def _texto_docx(data: bytes) -> str:
    import docx
    doc = docx.Document(io.BytesIO(data))
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for tabla in doc.tables:
        for fila in tabla.rows:
            celdas = [c.text.strip() for c in fila.cells if c.text.strip()]
            if celdas:
                partes.append(" | ".join(celdas))
    return "\n".join(partes)


def _texto_eml(data: bytes) -> str:
    import email
    from email import policy
    msg = email.message_from_bytes(data, policy=policy.default)
    cuerpo = ""
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                cuerpo += part.get_content()
    else:
        cuerpo = msg.get_content()
    asunto = msg.get("subject", "")
    de = msg.get("from", "")
    return f"Asunto: {asunto}\nDe: {de}\n\n{cuerpo}"


def extraer_texto(data: bytes, nombre: str = "", mime: str = "") -> tuple[str, str]:
    """Devuelve (texto, formato_detectado). Nunca lanza: ante error devuelve ('', ext)."""
    ext = _ext(nombre, mime)
    try:
        if ext == "pdf":
            return _texto_pdf(data), ext
        if ext in ("xlsx", "xls"):
            return _texto_xlsx(data), ext
        if ext in ("docx", "doc"):
            return _texto_docx(data), ext
        if ext == "eml":
            return _texto_eml(data), ext
        # csv / txt / desconocido → decodificar
        return data.decode("utf-8", errors="replace"), ext
    except Exception as e:
        log.error(f"extraer_texto({ext}): {e}")
        return "", ext


def descargar(url: str, timeout: int = 40) -> bytes:
    r = httpx.get(url, timeout=timeout, follow_redirects=True)
    r.raise_for_status()
    return r.content


# ─────────────────────────────────────────────────────────────
# 2. NORMALIZACIÓN (LLM chico → JSON estructurado)
# ─────────────────────────────────────────────────────────────
_NORM_SYS = (
    "Eres un extractor de datos de ÓRDENES DE COMPRA (purchase orders) de clientes industriales. "
    "Recibes el TEXTO crudo de una orden y devuelves SOLO un JSON válido, sin explicaciones ni ```.\n\n"
    "Esquema exacto:\n"
    "{\n"
    '  "cliente": "razón social del cliente que emite la orden (comprador), o \\"\\" si no está claro",\n'
    '  "po_number": "número de la orden de compra del cliente, tal cual aparece, o \\"\\"",\n'
    '  "moneda": "MXN | USD | \\"\\" si no se indica",\n'
    '  "items": [\n'
    '    {"part_number": "número de parte/modelo tal cual", "descripcion": "texto del renglón", '
    '"cantidad": number, "precio_unitario": number o null}\n'
    "  ],\n"
    '  "notas": "cualquier dato dudoso o que faltó, en una línea"\n'
    "}\n\n"
    "REGLAS:\n"
    "- NO inventes. Si un precio o cantidad no aparece, usa null (precio) o deja el item con lo que haya.\n"
    "- El part_number es el identificador del producto (SKU/modelo/mfr part no), NO la descripción.\n"
    "- Si el texto NO parece una orden de compra, devuelve items:[] y explica en notas.\n"
    "- Números sin separadores de miles ni símbolo de moneda (1234.56, no $1,234.56)."
)


def _parse_respuesta(raw: str) -> dict:
    """Extrae el JSON del esquema de la respuesta del modelo. Robusto ante ```-fences y texto extra:
    toma del primer '{' al último '}'. Lanza json.JSONDecodeError si no hay JSON válido."""
    raw = (raw or "").strip()
    i, j = raw.find("{"), raw.rfind("}")
    if i >= 0 and j > i:
        raw = raw[i:j + 1]
    datos = json.loads(raw)
    datos.setdefault("items", [])
    datos.setdefault("cliente", "")
    datos.setdefault("po_number", "")
    datos.setdefault("moneda", "")
    return datos


def normalizar(texto: str, model_id: str = "") -> dict:
    """Convierte el texto crudo del PO en datos estructurados. Devuelve el dict del esquema
    (con 'error' si algo falla, para que el chat lo reporte sin romperse)."""
    texto = (texto or "").strip()
    if not texto:
        return {"error": "documento vacío o ilegible", "items": []}
    # Limitar tamaño: un PO real cabe de sobra en ~15k chars; evita gasto en documentos enormes.
    recorte = texto[:15000]
    try:
        resp = llm.complete(
            model_id=model_id or PO_MODEL,
            system=_NORM_SYS,
            messages=[{"role": "user", "content": recorte}],
            max_tokens=2000,
            temperature=0,
        )
        return _parse_respuesta(resp.content[0].text)
    except json.JSONDecodeError as e:
        log.error(f"normalizar: JSON inválido: {e}")
        return {"error": "no pude estructurar la orden (JSON inválido del extractor)", "items": []}
    except Exception as e:
        log.error(f"normalizar: {e}")
        return {"error": f"error al normalizar: {e}", "items": []}


def _pdf_a_imagenes(data: bytes, max_pages: int = 3) -> list[bytes]:
    """Rasteriza las primeras páginas del PDF a PNG (para PDFs ESCANEADOS sin capa de texto).
    Cap de páginas para acotar el costo de visión."""
    import pypdfium2 as pdfium
    pdf = pdfium.PdfDocument(data)
    out: list[bytes] = []
    for i in range(min(len(pdf), max_pages)):
        pil = pdf[i].render(scale=2).to_pil()
        buf = io.BytesIO()
        pil.save(buf, format="PNG")
        out.append(buf.getvalue())
    return out


def normalizar_vision(imagenes: list[bytes], model_id: str = "") -> dict:
    """Fallback para PDFs ESCANEADOS: lee la orden desde imágenes con visión de Claude.
    Solo se usa cuando NO hay capa de texto (no encarece el caso normal)."""
    if not imagenes:
        return {"error": "sin imágenes que leer", "items": []}
    import base64
    import anthropic
    contenido: list = [{"type": "text", "text":
                        "Esta es una ORDEN DE COMPRA escaneada (imágenes). Extrae los datos al esquema pedido."}]
    for img in imagenes:
        contenido.append({"type": "image", "source": {
            "type": "base64", "media_type": "image/png",
            "data": base64.b64encode(img).decode()}})
    try:
        client = anthropic.Anthropic(api_key=os.environ.get("ANTHROPIC_API_KEY", ""))
        resp = client.messages.create(
            model=model_id or PO_MODEL, system=_NORM_SYS,
            max_tokens=2000, temperature=0,
            messages=[{"role": "user", "content": contenido}],
        )
        return _parse_respuesta(resp.content[0].text)
    except json.JSONDecodeError as e:
        log.error(f"normalizar_vision: JSON inválido: {e}")
        return {"error": "no pude estructurar la orden escaneada (JSON inválido)", "items": []}
    except Exception as e:
        log.error(f"normalizar_vision: {e}")
        return {"error": f"no pude leer el PDF escaneado con visión: {e}", "items": []}


def leer_po(url: str = "", data: bytes = b"", nombre: str = "", mime: str = "", model_id: str = "") -> dict:
    """Punto de entrada: baja el archivo (o usa data), extrae texto y lo normaliza.
    Devuelve {cliente, po_number, moneda, items[...], notas, formato} o {error}."""
    try:
        if not data:
            if not url:
                return {"error": "sin archivo ni URL", "items": []}
            data = descargar(url)
    except Exception as e:
        return {"error": f"no pude descargar el archivo: {e}", "items": []}

    texto, fmt = extraer_texto(data, nombre, mime)
    if not texto.strip():
        # PDF sin capa de texto = escaneado → fallback a visión (solo aquí, no en el caso normal).
        if fmt == "pdf":
            try:
                imgs = _pdf_a_imagenes(data)
            except Exception as e:
                imgs = []
                log.error(f"rasterizar PDF escaneado: {e}")
            if imgs:
                log.info(f"PDF sin texto (escaneado) → visión con {len(imgs)} página(s)")
                datos = normalizar_vision(imgs, model_id=model_id)
                datos["formato"] = "pdf(escaneado)"
                return datos
        return {"error": f"no pude extraer texto del archivo ({fmt})", "items": [], "formato": fmt}

    datos = normalizar(texto, model_id=model_id)
    datos["formato"] = fmt
    return datos
